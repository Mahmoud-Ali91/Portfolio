# extract_tables_by_table1_reset.py

import os
import re
import pandas as pd
from docx import Document

# === Setup ===
input_docx = "Comprehensive Aesthetic Services Market Analysis Tables.docx"
output_folder = "output_tables"
os.makedirs(output_folder, exist_ok=True)

doc = Document(input_docx)

def sanitize_filename(text):
    return re.sub(r'[^a-zA-Z0-9_]', '_', text.strip()).strip('_')[:60]

# === State ===
section_id_counter = 0
table_in_section = 0
current_section_id = None
last_detected_title = None

# Map paragraph elements
elements = list(doc.element.body)
paragraphs = doc.paragraphs
para_map = {}
p_idx = 0
for el in elements:
    if el.tag.endswith('}p') and p_idx < len(paragraphs):
        para_map[el] = paragraphs[p_idx]
        p_idx += 1

table_objects = doc.tables
table_obj_index = 0

# === Parse Document ===
for el in elements:
    tag = el.tag.split('}')[-1]

    if tag == 'p':
        para = para_map.get(el)
        text = para.text.strip()

        # Detect "Table X: Title"
        match = re.match(r"Table\s*(\d+)\s*:\s*(.+)", text, re.IGNORECASE)
        if match:
            table_number = int(match.group(1))
            title = sanitize_filename(match.group(2))

            if table_number == 1:
                section_id_counter += 1
                current_section_id = f"A{section_id_counter:02}"
                table_in_section = 1
            else:
                table_in_section += 1

            last_detected_title = title

    elif tag == 'tbl':
        table = table_objects[table_obj_index]
        table_obj_index += 1

        # Handle cases with no title
        if current_section_id is None:
            section_id_counter += 1
            current_section_id = f"A{section_id_counter:02}"
            table_in_section = 1

        elif last_detected_title is None:
            table_in_section += 1

        # Prepare filename
        title_part = last_detected_title if last_detected_title else "Untitled"
        filename = f"{current_section_id}_Table_{table_in_section}_{title_part}.csv"
        filepath = os.path.join(output_folder, filename)

        # Extract data
        data = []
        for row in table.rows:
            data.append([cell.text.strip() for cell in row.cells])

        try:
            df = pd.DataFrame(data[1:], columns=data[0])
        except Exception:
            df = pd.DataFrame(data)

        df.to_csv(filepath, index=False)
        print(f"✅ Saved: {filepath}")

        last_detected_title = None
